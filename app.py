import os, re, zipfile, tempfile, uuid, io, subprocess, shutil, json, datetime
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
import pandas as pd
from flask import Flask, request, jsonify, send_file, render_template_string
from docx import Document

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

OLLAMA_MODEL     = "llama3.2"
OLLAMA_URL       = "http://localhost:11434/api/generate"
OLLAMA_ENABLED   = False
OLLAMA_TIMEOUT   = 120
PARALLEL_WORKERS = 2

def _check_ollama():
    global OLLAMA_MODEL
    try:
        req   = urllib.request.urlopen("http://localhost:11434/api/tags", timeout=3)
        data  = json.loads(req.read())
        models = [m["name"].split(":")[0] for m in data.get("models", [])]
        for m in [OLLAMA_MODEL, "llama3.2", "mistral", "phi3"]:
            if m in models:
                OLLAMA_MODEL = m
                return True
        return False
    except Exception:
        return False

OLLAMA_ENABLED = _check_ollama()

LIBREOFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"
if not os.path.exists(LIBREOFFICE):
    _alt = os.path.expandvars(r"%LOCALAPPDATA%\Programs\LibreOffice\program\soffice.exe")
    LIBREOFFICE = _alt if os.path.exists(_alt) else (shutil.which("libreoffice") or shutil.which("soffice"))

# History stored in memory (persists while app is running)
HISTORY = []   # list of {id, name, date, count, errors, session_id}
SESSIONS = {}  # session_id → rows

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024

# ── History storage folder (saves Excel files to disk) ───────
HISTORY_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "history")
os.makedirs(HISTORY_DIR, exist_ok=True)

# ============================================================
# READ HELPERS
# ============================================================

def read_docx(path):
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t: lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells: lines.append(" | ".join(cells))
    return "\n".join(lines)

def convert_doc_to_docx(doc_path, out_dir):
    if not LIBREOFFICE: return None
    try:
        subprocess.run([LIBREOFFICE,"--headless","--convert-to","docx","--outdir",out_dir,doc_path],
                       capture_output=True, timeout=30)
        base = os.path.splitext(os.path.basename(doc_path))[0]
        conv = os.path.join(out_dir, base + ".docx")
        return conv if os.path.exists(conv) else None
    except Exception: return None

def read_pdf(path):
    if not PDF_SUPPORT: return ""
    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: lines.append(t)
    return "\n".join(lines)

def read_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx": return read_docx(path)
    elif ext == ".doc":
        tmp = tempfile.mkdtemp()
        conv = convert_doc_to_docx(path, tmp)
        return read_docx(conv) if conv else ""
    elif ext == ".pdf": return read_pdf(path)
    return ""

# ============================================================
# CLEAN HELPERS
# ============================================================

def clean_text(text):
    text = text.replace("\xa0", " ")
    return re.sub(r'[ \t]+', ' ', text)

def clean_amount(value):
    if not value: return ""
    value = str(value)
    value = re.sub(r'[,\s]', '', value)
    value = value.replace("/-", "")
    value = re.sub(r'-(?=\d{2}$)', '.', value)
    value = re.sub(r'(\.\d{2})\.\d{2}$', r'\1', value)
    value = value.strip()
    try: return "{:.2f}".format(float(value))
    except ValueError: return value

def is_blank(v):
    if v is None: return True
    return str(v).strip() in ("", "0", "0.0", "0.00", "None", "nan")

def extract_bill_no(filename):
    name = os.path.splitext(filename)[0]
    m = re.match(r'^(\d+)([A-Za-z]?)(?:[_\-\s]|$)', name)
    return (m.group(1) + m.group(2).upper()) if m else ""

COMPANY_MAP = {
    "GODIGIT":"Go Digit General Insurance Ltd","DIGIT":"Go Digit General Insurance Ltd",
    "UIIC":"United India Insurance Company Ltd","NICL":"National Insurance Company Ltd",
    "NIC":"National Insurance Company Ltd","NIA":"New India Assurance Company Ltd",
    "SBI":"SBI General Insurance Company Ltd","IFFCO":"Iffco-Tokio General Insurance Company Ltd",
    "ITGI":"Iffco-Tokio General Insurance Company Ltd","OIC":"The Oriental Insurance Company Ltd",
    "SGIC":"Shri Ram General Insurance Company Limited","SGI":"Shri Ram General Insurance Company Limited",
    "CHOLA":"Chola MS General Insurance Company Ltd","ACKO":"Acko General Insurance Limited",
    "ZUNO":"Zuno General Insurance Limited","RAHEJA":"Raheja QBE General Insurance Company Limited",
    "USGI":"Universal Sompo General Insurance Company Ltd","MGHDI":"Magma HDI General Insurance Company",
    "LIBERTY":"Liberty General Insurance Limited","HDFC":"HDFC ERGO General Insurance Company Ltd",
    "BAJAJ":"Bajaj Allianz General Insurance Company Ltd","TATA":"Tata AIG General Insurance Company Ltd",
    "ICICI":"ICICI Lombard General Insurance Company Ltd","RELIANCE":"Reliance General Insurance Company Ltd",
}

def detect_company(filename):
    up = filename.upper()
    for key, val in COMPANY_MAP.items():
        if key in up: return val
    return ""

# ============================================================
# EXTRACTION
# ============================================================

def extract_data(filename, text):
    bill_no = extract_bill_no(filename)
    company = detect_company(filename)

    date = ""
    for pat in [
        r'Invoice\s*Date\s*[:\-]?\s*([0-9]{1,2}[.\/\-][0-9]{1,2}[.\/\-][0-9]{2,4})',
        r'Bill\s*Date\s*[:\-]?\s*([0-9]{1,2}[.\/\-][0-9]{1,2}[.\/\-][0-9]{2,4})',
        r'Bill Date.*?\|\s*([0-9]{1,2}[.\/\-][0-9]{1,2}[.\/\-][0-9]{2,4})',
        r'Date\s*[:\-]?\s*([0-9]{1,2}[.\/\-][0-9]{1,2}[.\/\-][0-9]{2,4})',
        r'Dt\.?\s*[:\-]?\s*([0-9]{1,2}[.\/\-][0-9]{1,2}[.\/\-][0-9]{2,4})',
    ]:
        ms = re.findall(pat, text, re.IGNORECASE)
        if ms: date = ms[0]; break

    cgst = ""
    for pat in [
        r'Service\s*CGST\s*Amount\s*\|[^\d\n]*([0-9,]+\.[0-9]{2})',
        r'Service\s*CGST\s*Amount[^\d\n]{0,20}([0-9,]+\.[0-9]{2})',
        r'CGST@0?9%\s*\|[^\d\n]*([0-9,]+\.[0-9]{2})',
        r'CGST@0?9%[^\d\n]{0,15}([0-9,]+\.[0-9]{2})',
        r'C\s*GST\s*@\s*0?9\s*%[^\d\n]{0,80}([0-9,]+\.[0-9]{2})',
        r'CGST\s*@\s*0?9\s*%[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'CGST[^\d\n]{0,25}([0-9,]+\.[0-9]{2})',
    ]:
        ms = re.findall(pat, text, re.IGNORECASE)
        if ms: cgst = ms[-1]; break

    sgst = ""
    for pat in [
        r'Service\s*SGST\s*Amount\s*\|[^\d\n]*([0-9,]+\.[0-9]{2})',
        r'Service\s*SGST\s*Amount[^\d\n]{0,20}([0-9,]+\.[0-9]{2})',
        r'SCGST@0?9%\s*\|[^\d\n]*([0-9,]+\.[0-9]{2})',
        r'SCGST@0?9%[^\d\n]{0,15}([0-9,]+\.[0-9]{2})',
        r'S\s*GST\s*@\s*0?9\s*%[^\d\n]{0,80}([0-9,]+\.[0-9]{2})',
        r'SGST\s*@\s*0?9\s*%[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'SGST[^\d\n]{0,25}([0-9,]+\.[0-9]{2})',
    ]:
        ms = re.findall(pat, text, re.IGNORECASE)
        if ms: sgst = ms[-1]; break

    igst = ""
    for pat in [
        r'IGST\s*Amount\s*\|[^\d\n]*([0-9,]+\.[0-9]{2})',
        r'IGST\s*Amount[^\d\n]{0,20}([0-9,]+\.[0-9]{2})',
        r'ISGST\s*@\s*18\s*%[^\d\n]{0,80}([0-9,]+\.[0-9]{2})',
        r'ICGST@\s*18\s*%[^\d\n]{0,80}([0-9,]+\.[0-9]{2})',
        r'ICGST@18%[^\d\n]{0,80}([0-9,]+\.[0-9]{2})',
        r'GST\s*Tax\s*\|[^\d\n]*@18[^\d\n]*([0-9,]+\.[0-9]{2})',
        r'GST\s*Tax[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'IGST\s*@\s*18\s*%[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'IGST[^\d\n]{0,25}([0-9,]+\.[0-9]{2})',
    ]:
        ms = re.findall(pat, text, re.IGNORECASE)
        if ms: igst = ms[-1]; break

    grand_total = ""
    for pat in [
        r'Grand\s*Total[^\n]*?([0-9,]+\.[0-9]{2}(?:\.[0-9]{2})?)\s*(?:\|.*)?$',
        r'Grand\s*Total[^\n]*?([0-9,]+\-[0-9]{2})\s*(?:\|.*)?$',
        r'Total\s*Amount\s*\(?In\s*Fig[^\n]*?Rs\.?\s*([0-9,]+(?:\.[0-9]{2})?)',
        r'Total\s*Amount\s*\(?In\s*Fig[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'Total\s*Amount\s*(?:Payable)?[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'Or\s*Say\s*Rs\.?\s*([0-9,]+(?:\.[0-9]{2})?)',
        r'Invoice\s*(?:Value|Amount)[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'Net\s*(?:Payable|Amount)[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
        r'Amount\s*Payable[^\d\n]{0,30}([0-9,]+\.[0-9]{2})',
    ]:
        ms = re.findall(pat, text, re.IGNORECASE | re.MULTILINE)
        if ms: grand_total = ms[-1]; break

    return {
        "Bill No":     bill_no,
        "Company":     company,
        "Date":        date,
        "Grand Total": clean_amount(grand_total),
        "IGST":        clean_amount(igst),
        "CGST":        clean_amount(cgst),
        "SGST":        clean_amount(sgst),
        "_ai_filled":  False,
    }

# ============================================================
# OLLAMA FALLBACK
# ============================================================

def _ollama_ask(prompt):
    payload = json.dumps({
        "model": OLLAMA_MODEL, "prompt": prompt, "stream": False,
        "options": {"temperature": 0, "num_predict": 150, "num_ctx": 2048}
    }).encode("utf-8")
    req = urllib.request.Request(OLLAMA_URL, data=payload,
                                  headers={"Content-Type":"application/json"}, method="POST")
    with urllib.request.urlopen(req, timeout=OLLAMA_TIMEOUT) as resp:
        return json.loads(resp.read()).get("response","").strip()

def ollama_fill_missing(filename, text, row):
    if not OLLAMA_ENABLED: return row
    igst_f = not is_blank(row.get("IGST"))
    cgst_f = not is_blank(row.get("CGST"))
    sgst_f = not is_blank(row.get("SGST"))
    missing = []
    if is_blank(row.get("Date")): missing.append("invoice_date")
    if is_blank(row.get("Grand Total")): missing.append("grand_total")
    if not igst_f and not cgst_f and not sgst_f:
        missing += ["igst","cgst","sgst"]
    if total_f := not is_blank(row.get("Grand Total")):
        try:
            t = float(row.get("Grand Total","0") or "0")
            g = sum(float(row.get(k,"0") or "0") for k in ("CGST","SGST","IGST"))
            if t < 50 or (g > 0 and g > t * 1.01):
                if "grand_total" not in missing: missing.append("grand_total")
        except: pass
    if not missing: return row
    prompt = f"""Extract fields {', '.join(missing)} from invoice text.
Rules: grand_total=final payable incl taxes, invoice_date=as written e.g. 12.09.2025, cgst/sgst/igst=plain numbers.
Return ONLY JSON, no explanation.
Invoice: {filename}
Text: {text[:3000]}
JSON:"""
    try:
        raw  = _ollama_ask(prompt)
        raw  = re.sub(r'```(?:json)?','',raw).strip().rstrip('`').strip()
        m    = re.search(r'\{.*\}', raw, re.DOTALL)
        if not m: return row
        data = json.loads(m.group())
        field_map = {"invoice_date":("Date",False),"grand_total":("Grand Total",True),
                     "cgst":("CGST",True),"sgst":("SGST",True),"igst":("IGST",True)}
        helped = False
        for ak,(ek,ia) in field_map.items():
            if ak not in missing: continue
            v = str(data.get(ak,"")).strip()
            if v:
                row[ek] = clean_amount(v) if ia else v
                helped  = True
        if helped: row["_ai_filled"] = True
    except Exception as e:
        print(f"  AI error [{filename}]: {str(e)[:80]}")
    return row

# ============================================================
# PROCESS
# ============================================================

def process_one(fname, fpath):
    try:
        raw  = read_file(fpath)
        text = clean_text(raw)
        if not text.strip():
            return None, {"File": fname, "Error": "Could not read file. For .doc files, LibreOffice must be installed."}
        row = extract_data(fname, text)
        row = ollama_fill_missing(fname, text, row)
        return row, None
    except Exception as e:
        return None, {"File": fname, "Error": str(e)}

def process_zip(zip_path):
    results, errors, seen_nos = [], [], {}
    with tempfile.TemporaryDirectory() as tmp:
        with zipfile.ZipFile(zip_path,'r') as zf: zf.extractall(tmp)
        all_files = []
        for root,_,files in os.walk(tmp):
            for fname in files:
                if fname.startswith("~$"): continue
                ext = os.path.splitext(fname)[1].lower()
                if ext in (".docx",".doc",".pdf"):
                    all_files.append((fname, os.path.join(root,fname), ext))
        all_files.sort(key=lambda x:(x[0], 0 if x[2]==".docx" else 1))
        seen_files, unique = set(), []
        for fname,fpath,ext in all_files:
            base = os.path.splitext(fname)[0]
            if base not in seen_files:
                seen_files.add(base); unique.append((fname,fpath))
        with ThreadPoolExecutor(max_workers=PARALLEL_WORKERS) as ex:
            futures = {ex.submit(process_one,fn,fp):fn for fn,fp in unique}
            for future in as_completed(futures):
                row,err = future.result()
                if err: errors.append(err)
                elif row:
                    bn = row["Bill No"]
                    if bn and bn in seen_nos:
                        ex_ = results[seen_nos[bn]]
                        if sum(1 for v in row.values() if v) > sum(1 for v in ex_.values() if v):
                            results[seen_nos[bn]] = row
                    else:
                        seen_nos[bn] = len(results); results.append(row)
    def sk(r):
        m = re.match(r'(\d+)', r.get("Bill No",""))
        return (int(m.group(1)) if m else 9999999, r.get("Bill No",""))
    results.sort(key=sk)
    return results, errors

def make_excel(rows):
    cols   = ["Bill No","Company","Date","Grand Total","IGST","CGST","SGST"]
    export = [{k:v for k,v in r.items() if not k.startswith("_")} for r in rows]
    df     = pd.DataFrame(export, columns=cols)
    buf    = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="GST Details")
        ws = writer.sheets["GST Details"]
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        # Header styling
        hdr_fill = PatternFill("solid", fgColor="1E3A5F")
        hdr_font = Font(bold=True, color="FFFFFF", size=11)
        for cell in ws[1]:
            cell.fill = hdr_fill; cell.font = hdr_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 25
        # Data rows alternating colour
        light = PatternFill("solid", fgColor="F0F4FA")
        for i,row in enumerate(ws.iter_rows(min_row=2, max_row=ws.max_row), start=2):
            for cell in row:
                cell.alignment = Alignment(horizontal="center")
                if i % 2 == 0: cell.fill = light
        # Column widths
        for col in ws.columns:
            mx = max(len(str(c.value or "")) for c in col)
            ws.column_dimensions[col[0].column_letter].width = min(mx+6, 45)
    buf.seek(0)
    return buf

# ============================================================
# HTML
# ============================================================

HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1.0"/>
<title>GST Invoice Manager</title>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;600;700&family=Mulish:wght@300;400;500;600;700&display=swap" rel="stylesheet"/>
<style>
:root{
  --navy:#1E3A5F;--navy2:#2B4F82;--gold:#C9A84C;--gold2:#E8C96A;
  --bg:#F5F7FA;--white:#FFFFFF;--text:#1A2B3C;--muted:#6B7E94;
  --border:#D8E2EE;--success:#2E7D52;--danger:#C0392B;--warn:#C87A1A;
  --shadow:0 2px 12px rgba(30,58,95,.10);--shadow2:0 8px 32px rgba(30,58,95,.16);
}
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:'Mulish',sans-serif;background:var(--bg);color:var(--text);min-height:100vh;}

/* ── NAV ── */
nav{
  background:var(--navy);height:64px;display:flex;align-items:center;
  padding:0 2rem;gap:0;position:sticky;top:0;z-index:100;
  box-shadow:0 2px 16px rgba(0,0,0,.18);
}
.nav-brand{
  font-family:'Playfair Display',serif;font-size:1.35rem;font-weight:700;
  color:var(--white);letter-spacing:.02em;margin-right:3rem;white-space:nowrap;
}
.nav-brand span{color:var(--gold);}
.nav-links{display:flex;gap:.25rem;flex:1;}
.nav-link{
  color:rgba(255,255,255,.7);font-size:.82rem;font-weight:600;letter-spacing:.06em;
  text-transform:uppercase;padding:.5rem 1.1rem;border-radius:6px;cursor:pointer;
  transition:all .2s;border:none;background:none;
}
.nav-link:hover{color:var(--white);background:rgba(255,255,255,.1);}
.nav-link.active{color:var(--gold);background:rgba(201,168,76,.12);}
.nav-status{
  font-size:.72rem;color:rgba(255,255,255,.5);padding:.3rem .8rem;
  border:1px solid rgba(255,255,255,.15);border-radius:20px;white-space:nowrap;
}

/* ── PAGES ── */
.page{display:none;min-height:calc(100vh - 64px);padding:2.5rem 2rem;}
.page.active{display:block;}
.page-inner{max-width:1100px;margin:0 auto;}

/* ── PAGE HEADER ── */
.page-header{margin-bottom:2rem;}
.page-header h1{font-family:'Playfair Display',serif;font-size:1.8rem;font-weight:700;color:var(--navy);margin-bottom:.4rem;}
.page-header p{color:var(--muted);font-size:.9rem;}

/* ── CARDS ── */
.card{background:var(--white);border-radius:12px;border:1px solid var(--border);padding:2rem;box-shadow:var(--shadow);}
.card+.card{margin-top:1.5rem;}

/* ── UPLOAD ZONE ── */
#drop-zone{
  border:2px dashed var(--border);border-radius:12px;padding:3.5rem 2rem;
  text-align:center;cursor:pointer;transition:all .25s;background:var(--bg);
  position:relative;
}
#drop-zone:hover,#drop-zone.drag{border-color:var(--gold);background:#FDFAF3;}
#drop-zone input{position:absolute;inset:0;opacity:0;cursor:pointer;}
.upload-icon{width:56px;height:56px;border-radius:14px;background:var(--navy);
  display:flex;align-items:center;justify-content:center;margin:0 auto 1.2rem;}
.upload-icon svg{width:28px;height:28px;fill:var(--white);}
#drop-zone h3{font-family:'Playfair Display',serif;font-size:1.2rem;color:var(--navy);margin-bottom:.5rem;}
#drop-zone p{color:var(--muted);font-size:.83rem;}

/* ── FILE CHIP ── */
#file-chip{
  display:none;align-items:center;gap:.75rem;margin-top:1.25rem;
  background:#EEF4FB;border:1px solid var(--border);border-radius:8px;
  padding:.7rem 1rem;font-size:.83rem;
}
#file-chip.show{display:flex;}
.chip-icon{width:32px;height:32px;background:var(--navy);border-radius:6px;
  display:flex;align-items:center;justify-content:center;flex-shrink:0;}
.chip-icon svg{width:16px;height:16px;fill:white;}
.chip-name{font-weight:600;color:var(--navy);}
.chip-size{color:var(--muted);font-size:.76rem;margin-left:.25rem;}

/* ── BUTTONS ── */
.btn{
  display:inline-flex;align-items:center;gap:.5rem;font-family:'Mulish',sans-serif;
  font-weight:700;font-size:.83rem;letter-spacing:.04em;border:none;border-radius:8px;
  padding:.7rem 1.6rem;cursor:pointer;transition:all .2s;white-space:nowrap;
}
.btn-primary{background:var(--navy);color:var(--white);}
.btn-primary:hover{background:var(--navy2);transform:translateY(-1px);box-shadow:var(--shadow2);}
.btn-primary:disabled{opacity:.4;cursor:not-allowed;transform:none;box-shadow:none;}
.btn-gold{background:var(--gold);color:var(--navy);}
.btn-gold:hover{background:var(--gold2);transform:translateY(-1px);box-shadow:var(--shadow2);}
.btn-outline{background:transparent;color:var(--navy);border:1.5px solid var(--border);}
.btn-outline:hover{border-color:var(--navy);background:#F0F4FA;}
.btn-danger-sm{background:transparent;color:var(--danger);border:1px solid var(--danger);
  padding:.35rem .8rem;font-size:.75rem;border-radius:6px;}
.btn-danger-sm:hover{background:var(--danger);color:white;}

/* ── PROGRESS ── */
#prog-wrap{display:none;margin-top:1.5rem;}
#prog-wrap.show{display:block;}
.prog-track{height:8px;background:var(--border);border-radius:999px;overflow:hidden;margin:.5rem 0;}
.prog-fill{height:100%;width:0%;background:linear-gradient(90deg,var(--navy),var(--gold));
  border-radius:999px;transition:width .5s ease;}
.prog-info{display:flex;justify-content:space-between;font-size:.78rem;color:var(--muted);}

/* ── STATS ROW ── */
.stats-row{display:grid;grid-template-columns:repeat(4,1fr);gap:1rem;margin-bottom:1.5rem;}
.stat-box{background:var(--white);border:1px solid var(--border);border-radius:10px;
  padding:1.1rem 1.25rem;box-shadow:var(--shadow);transition:box-shadow .2s;}
.stat-box:hover{box-shadow:var(--shadow2);}
.stat-num{font-family:'Playfair Display',serif;font-size:1.8rem;font-weight:700;color:var(--navy);line-height:1;}
.stat-lbl{font-size:.72rem;color:var(--muted);margin-top:.3rem;font-weight:600;text-transform:uppercase;letter-spacing:.05em;}
.stat-box.gold .stat-num{color:var(--gold);}
.stat-box.red .stat-num{color:var(--danger);}
.stat-box.green .stat-num{color:var(--success);}

/* ── TABLE ── */
.tbl-wrap{overflow-x:auto;border-radius:10px;border:1px solid var(--border);}
table{width:100%;border-collapse:collapse;font-size:.82rem;}
thead{background:var(--navy);}
th{
  padding:.8rem 1rem;text-align:left;color:rgba(255,255,255,.85);font-weight:600;
  font-size:.72rem;letter-spacing:.07em;text-transform:uppercase;white-space:nowrap;
}
th:first-child{border-radius:10px 0 0 0;}th:last-child{border-radius:0 10px 0 0;}
td{padding:.7rem 1rem;border-bottom:1px solid var(--border);vertical-align:middle;color:var(--text);}
tr:last-child td{border-bottom:none;}
tbody tr{transition:background .15s;}
tbody tr:hover td{background:#EEF4FB;}
.td-bill{font-weight:700;color:var(--navy);font-size:.88rem;}
.td-company{color:var(--text);max-width:200px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;}
.td-date{color:var(--muted);white-space:nowrap;}
.td-amt{font-weight:600;color:var(--navy);}
.td-gst{color:var(--success);}
.td-na{color:var(--border);font-size:.75rem;}

/* ── ERRORS ── */
.error-box{background:#FEF0EF;border:1px solid #F5C6C3;border-radius:10px;
  padding:1.25rem;margin-top:1.25rem;}
.error-box h4{color:var(--danger);font-size:.88rem;font-weight:700;margin-bottom:.75rem;}
.error-item{font-size:.78rem;color:#7A2828;padding:.3rem 0;border-bottom:1px solid #F5C6C3;}
.error-item:last-child{border-bottom:none;}
.error-item strong{color:var(--danger);}

/* ── ACTIONS ── */
.actions-row{display:flex;gap:.75rem;margin-top:1.5rem;flex-wrap:wrap;}

/* ── HISTORY ── */
.history-empty{text-align:center;padding:3rem;color:var(--muted);}
.history-empty svg{width:48px;height:48px;fill:var(--border);display:block;margin:0 auto 1rem;}
.history-table{width:100%;border-collapse:collapse;font-size:.83rem;}
.history-table th{
  padding:.7rem 1rem;text-align:left;background:var(--navy);color:rgba(255,255,255,.85);
  font-size:.72rem;letter-spacing:.07em;text-transform:uppercase;font-weight:600;
}
.history-table th:first-child{border-radius:10px 0 0 0;}
.history-table th:last-child{border-radius:0 10px 0 0;}
.history-table td{padding:.7rem 1rem;border-bottom:1px solid var(--border);vertical-align:middle;}
.history-table tr:last-child td{border-bottom:none;}
.history-table tbody tr{transition:background .15s;}
.history-table tbody tr:hover td{background:#EEF4FB;}
.h-badge{display:inline-block;padding:.2rem .6rem;border-radius:20px;font-size:.72rem;font-weight:600;}
.h-badge.ok{background:#E8F5EE;color:var(--success);}
.h-badge.warn{background:#FEF3E2;color:var(--warn);}

/* ── HELP ── */
.help-grid{display:grid;grid-template-columns:1fr 1fr;gap:1.25rem;}
.help-card{background:var(--white);border:1px solid var(--border);border-radius:10px;
  padding:1.5rem;box-shadow:var(--shadow);transition:all .2s;}
.help-card:hover{border-color:var(--gold);box-shadow:var(--shadow2);transform:translateY(-2px);}
.help-card h3{font-family:'Playfair Display',serif;font-size:1rem;color:var(--navy);
  margin-bottom:.6rem;display:flex;align-items:center;gap:.5rem;}
.help-card p,.help-card li{font-size:.83rem;color:var(--muted);line-height:1.6;}
.help-card ul{padding-left:1.1rem;}
.step-num{width:24px;height:24px;background:var(--navy);color:white;border-radius:50%;
  display:inline-flex;align-items:center;justify-content:center;font-size:.72rem;
  font-weight:700;flex-shrink:0;}
.faq-item{border-bottom:1px solid var(--border);padding:1rem 0;}
.faq-item:last-child{border-bottom:none;}
.faq-q{font-weight:700;color:var(--navy);font-size:.88rem;margin-bottom:.35rem;}
.faq-a{font-size:.82rem;color:var(--muted);line-height:1.6;}
.highlight-box{background:linear-gradient(135deg,#EEF4FB,#FDF8EC);border:1px solid var(--border);
  border-radius:10px;padding:1.25rem 1.5rem;margin-top:1.25rem;}
.highlight-box h4{color:var(--navy);font-size:.88rem;font-weight:700;margin-bottom:.5rem;}
.highlight-box p{font-size:.82rem;color:var(--muted);line-height:1.6;}

/* ── RESPONSIVE ── */
@media(max-width:768px){
  .stats-row{grid-template-columns:1fr 1fr;}
  .help-grid{grid-template-columns:1fr;}
  .nav-link{padding:.5rem .7rem;font-size:.75rem;}
  .nav-brand{margin-right:1rem;font-size:1.1rem;}
}
</style>
</head>
<body>

<nav>
  <div class="nav-brand">GST<span>.</span>Manager</div>
  <div class="nav-links">
    <button class="nav-link active" onclick="showPage('extract')">Extract Bills</button>
    <button class="nav-link" onclick="showPage('history')">History</button>
    <button class="nav-link" onclick="showPage('help')">Help Guide</button>
  </div>
  <div class="nav-status" id="nav-status">Ready</div>
</nav>

<!-- =========================================================
     PAGE: EXTRACT
========================================================= -->
<div class="page active" id="page-extract">
<div class="page-inner">
  <div class="page-header">
    <h1>Extract GST Details</h1>
    <p>Upload a ZIP file containing your invoice bills to extract GST information automatically.</p>
  </div>

  <div class="card">
    <div id="drop-zone">
      <input type="file" id="file-input" accept=".zip"/>
      <div class="upload-icon">
        <svg viewBox="0 0 24 24"><path d="M19.35 10.04A7.49 7.49 0 0 0 12 4C9.11 4 6.6 5.64 5.35 8.04A5.994 5.994 0 0 0 0 14c0 3.31 2.69 6 6 6h13c2.76 0 5-2.24 5-5 0-2.64-2.05-4.78-4.65-4.96zM14 13v4h-4v-4H7l5-5 5 5h-3z"/></svg>
      </div>
      <h3>Drop your ZIP file here</h3>
      <p>Or click to browse &nbsp;&middot;&nbsp; Supports .docx, .doc, .pdf &nbsp;&middot;&nbsp; Up to 500+ bills per batch</p>
    </div>

    <div id="file-chip">
      <div class="chip-icon">
        <svg viewBox="0 0 24 24"><path d="M6 2c-1.1 0-2 .9-2 2v16c0 1.1.9 2 2 2h12c1.1 0 2-.9 2-2V8l-6-6H6zm7 7V3.5L18.5 9H13z"/></svg>
      </div>
      <div>
        <span class="chip-name" id="chip-name"></span>
        <span class="chip-size" id="chip-size"></span>
      </div>
      <button class="btn btn-outline" style="margin-left:auto;padding:.35rem .8rem;font-size:.75rem" onclick="clearFile()">Remove</button>
    </div>

    <div style="margin-top:1.25rem;display:flex;gap:.75rem;align-items:center;">
      <button class="btn btn-primary" id="process-btn" disabled onclick="processFile()">
        <svg width="16" height="16" viewBox="0 0 24 24" fill="white"><path d="M8 5v14l11-7z"/></svg>
        Process Bills
      </button>
      <span id="hint-text" style="font-size:.78rem;color:var(--muted)">Select a ZIP file to begin</span>
    </div>

    <div id="prog-wrap">
      <div class="prog-info"><span id="prog-msg">Preparing...</span><span id="prog-pct">0%</span></div>
      <div class="prog-track"><div class="prog-fill" id="prog-fill"></div></div>
    </div>
  </div>

  <!-- RESULTS -->
  <div id="results-section" style="display:none;margin-top:1.5rem;">
    <div class="stats-row">
      <div class="stat-box">
        <div class="stat-num" id="s-total">0</div>
        <div class="stat-lbl">Total Bills</div>
      </div>
      <div class="stat-box green">
        <div class="stat-num" id="s-ok">0</div>
        <div class="stat-lbl">Processed</div>
      </div>
      <div class="stat-box red">
        <div class="stat-num" id="s-err">0</div>
        <div class="stat-lbl">Errors</div>
      </div>
      <div class="stat-box gold">
        <div class="stat-num" id="s-blank">0</div>
        <div class="stat-lbl">Incomplete Fields</div>
      </div>
    </div>

    <div class="card" style="padding:0;overflow:hidden;">
      <div style="padding:1rem 1.5rem;border-bottom:1px solid var(--border);display:flex;align-items:center;justify-content:space-between;">
        <div>
          <div style="font-family:'Playfair Display',serif;font-size:1rem;font-weight:600;color:var(--navy)">Extracted Data</div>
          <div style="font-size:.76rem;color:var(--muted);margin-top:.15rem">Review the results below before downloading</div>
        </div>
        <div style="display:flex;gap:.5rem;">
          <input type="text" id="search-box" placeholder="Search bill or company..." oninput="filterTable()"
            style="padding:.45rem .85rem;border:1px solid var(--border);border-radius:6px;font-size:.8rem;font-family:Mulish,sans-serif;color:var(--text);outline:none;width:220px;">
        </div>
      </div>
      <div class="tbl-wrap">
        <table>
          <thead>
            <tr>
              <th>Bill No</th><th>Company</th><th>Date</th>
              <th>Grand Total</th><th>IGST</th><th>CGST</th><th>SGST</th>
            </tr>
          </thead>
          <tbody id="tbl-body"></tbody>
        </table>
      </div>
    </div>

    <div id="error-box" style="display:none;" class="error-box">
      <h4>Files that could not be read</h4>
      <div id="error-list"></div>
    </div>

    <div class="actions-row">
      <button class="btn btn-gold" onclick="downloadExcel()">
        <svg width="16" height="16" viewBox="0 0 24 24" fill="currentColor"><path d="M19 9h-4V3H9v6H5l7 7 7-7zM5 18v2h14v-2H5z"/></svg>
        Download Excel
      </button>
      <button class="btn btn-outline" onclick="resetAll()">New Extraction</button>
    </div>
  </div>

</div>
</div>

<!-- =========================================================
     PAGE: HISTORY
========================================================= -->
<div class="page" id="page-history">
<div class="page-inner">
  <div class="page-header">
    <h1>Extraction History</h1>
    <p>All previous extractions from this session are saved here. Download any previous Excel file anytime.</p>
  </div>
  <div class="card" style="padding:0;overflow:hidden;">
    <div style="padding:1rem 1.5rem;border-bottom:1px solid var(--border);display:flex;align-items:center;justify-content:space-between;">
      <div style="font-family:'Playfair Display',serif;font-size:1rem;font-weight:600;color:var(--navy)">Previous Extractions</div>
      <button class="btn btn-outline" style="padding:.4rem .9rem;font-size:.76rem" onclick="clearHistory()">Clear All</button>
    </div>
    <div id="history-body"></div>
  </div>
</div>
</div>

<!-- =========================================================
     PAGE: HELP
========================================================= -->
<div class="page" id="page-help">
<div class="page-inner">
  <div class="page-header">
    <h1>Help Guide</h1>
    <p>Everything you need to know about using GST Manager effectively.</p>
  </div>

  <div class="help-grid">
    <div class="help-card">
      <h3><span class="step-num">1</span> Prepare Your Bills</h3>
      <ul>
        <li>Collect all your invoice files (.docx, .doc, or .pdf)</li>
        <li>Put them all into one folder</li>
        <li>Select all files, right-click and choose "Send to &rarr; Compressed (ZIP) folder"</li>
        <li>You can include 500+ bills in a single ZIP</li>
      </ul>
    </div>
    <div class="help-card">
      <h3><span class="step-num">2</span> Upload and Process</h3>
      <ul>
        <li>Go to "Extract Bills" page</li>
        <li>Drag and drop your ZIP file, or click to browse</li>
        <li>Click "Process Bills" and wait for extraction to complete</li>
        <li>Larger batches may take a few minutes</li>
      </ul>
    </div>
    <div class="help-card">
      <h3><span class="step-num">3</span> Review Results</h3>
      <ul>
        <li>Check the extracted data in the table</li>
        <li>Use the search box to find a specific bill or company</li>
        <li>Blank (—) in IGST means the bill is intra-state (CGST+SGST used instead) — this is correct</li>
        <li>Blank (—) in CGST/SGST means the bill is inter-state (IGST used) — also correct</li>
      </ul>
    </div>
    <div class="help-card">
      <h3><span class="step-num">4</span> Download Excel</h3>
      <ul>
        <li>Click "Download Excel" to save the extracted data</li>
        <li>The Excel file is automatically formatted with headers</li>
        <li>All previous extractions are saved in the History tab</li>
        <li>You can re-download any previous extraction anytime</li>
      </ul>
    </div>
  </div>

  <div class="highlight-box" style="margin-top:1.25rem;">
    <h4>Understanding GST Types</h4>
    <p>Indian GST has two types depending on whether the transaction is within the same state or across states. <strong>IGST</strong> (Integrated GST at 18%) applies to inter-state transactions. <strong>CGST + SGST</strong> (Central and State GST at 9% each) apply to intra-state transactions. A single bill will always use one or the other — never both. So if a bill shows IGST, the CGST and SGST columns will be blank — that is correct and expected.</p>
  </div>

  <div class="card" style="margin-top:1.25rem;">
    <div style="font-family:'Playfair Display',serif;font-size:1.1rem;font-weight:600;color:var(--navy);margin-bottom:1rem;">Frequently Asked Questions</div>

    <div class="faq-item">
      <div class="faq-q">Why are some fields showing a dash (—)?</div>
      <div class="faq-a">A dash means the field was not found in the bill. This can be because the bill uses a different format, the file could not be read (especially for older .doc files), or the field genuinely does not exist on that bill.</div>
    </div>
    <div class="faq-item">
      <div class="faq-q">Why are .doc files showing as errors?</div>
      <div class="faq-a">Older .doc format files require LibreOffice to be installed on your computer. Download it free from libreoffice.org, install it, and restart the app. After that, .doc files will be read automatically.</div>
    </div>
    <div class="faq-item">
      <div class="faq-q">Can I process files from multiple folders?</div>
      <div class="faq-a">Yes. Just put all your bill files into one ZIP regardless of their original folder structure. The app will scan all files inside the ZIP automatically.</div>
    </div>
    <div class="faq-item">
      <div class="faq-q">How many bills can I process at once?</div>
      <div class="faq-a">There is no hard limit. The app has been tested with 500+ bills in a single batch. Larger batches simply take more time to process.</div>
    </div>
    <div class="faq-item">
      <div class="faq-q">Is my data sent anywhere?</div>
      <div class="faq-a">No. This application runs entirely on your computer. Your bills and extracted data never leave your machine.</div>
    </div>
    <div class="faq-item">
      <div class="faq-q">Why does the Bill No show wrong sometimes?</div>
      <div class="faq-a">The Bill Number is read from the filename. Name your files starting with the bill number, for example: 319A_NIA_Claim.docx. The app reads the number and optional letter at the start of the filename.</div>
    </div>
  </div>
</div>
</div>

<script>
let sessionId = null;
let allRows   = [];
let history   = [];

// ── NAV ──────────────────────────────────────────────────────
function showPage(id) {
  document.querySelectorAll('.page').forEach(p => p.classList.remove('active'));
  document.querySelectorAll('.nav-link').forEach(l => l.classList.remove('active'));
  document.getElementById('page-'+id).classList.add('active');
  event.currentTarget.classList.add('active');
  if (id === 'history') renderHistory();
}

// ── UPLOAD ───────────────────────────────────────────────────
const dropZone = document.getElementById('drop-zone');
const fileInput = document.getElementById('file-input');

dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('drag'); });
dropZone.addEventListener('dragleave', () => dropZone.classList.remove('drag'));
dropZone.addEventListener('drop', e => {
  e.preventDefault(); dropZone.classList.remove('drag');
  const f = e.dataTransfer.files[0];
  if (f && f.name.endsWith('.zip')) setFile(f);
  else alert('Please drop a .zip file.');
});
fileInput.addEventListener('change', () => { if (fileInput.files[0]) setFile(fileInput.files[0]); });

function setFile(f) {
  document.getElementById('chip-name').textContent = f.name;
  document.getElementById('chip-size').textContent = '(' + (f.size/1024/1024).toFixed(1) + ' MB)';
  document.getElementById('file-chip').classList.add('show');
  document.getElementById('process-btn').disabled = false;
  document.getElementById('hint-text').textContent = 'Ready to process';
  document.getElementById('process-btn')._file = f;
  document.getElementById('process-btn')._name = f.name;
}
function clearFile() {
  fileInput.value = '';
  document.getElementById('file-chip').classList.remove('show');
  document.getElementById('process-btn').disabled = true;
  document.getElementById('hint-text').textContent = 'Select a ZIP file to begin';
}

// ── PROGRESS ─────────────────────────────────────────────────
let progInterval = null;
function startProgress() {
  let pct = 0;
  document.getElementById('prog-wrap').classList.add('show');
  const msgs = ['Reading ZIP file...','Scanning bill documents...','Extracting invoice data...','Processing GST fields...','Sorting results...','Finalising...'];
  let mi = 0;
  progInterval = setInterval(() => {
    pct = Math.min(pct + Math.random()*1.8, 90);
    document.getElementById('prog-fill').style.width = pct+'%';
    document.getElementById('prog-pct').textContent = Math.floor(pct)+'%';
    if (mi < msgs.length && pct > mi*15) document.getElementById('prog-msg').textContent = msgs[mi++];
  }, 400);
}
function endProgress() {
  clearInterval(progInterval);
  document.getElementById('prog-fill').style.width = '100%';
  document.getElementById('prog-pct').textContent = '100%';
  document.getElementById('prog-msg').textContent = 'Complete';
}

// ── PROCESS ──────────────────────────────────────────────────
async function processFile() {
  const btn = document.getElementById('process-btn');
  const f   = btn._file;
  if (!f) return;
  btn.disabled = true;
  document.getElementById('results-section').style.display = 'none';
  document.getElementById('nav-status').textContent = 'Processing...';
  startProgress();
  const form = new FormData();
  form.append('file', f);
  try {
    const res  = await fetch('/upload', { method:'POST', body:form });
    const data = await res.json();
    endProgress();
    if (data.error) { alert('Error: ' + data.error); return; }
    sessionId = data.session_id;
    allRows   = data.results;
    document.getElementById('nav-status').textContent = data.results.length + ' bills processed';
    renderResults(data.results, data.errors);
    // Save to history
    addToHistory(btn._name, data.results.length, data.errors.length, data.session_id);
  } catch(err) {
    endProgress();
    alert('Could not connect to server. Is the app running?');
  } finally {
    btn.disabled = false;
  }
}

// ── RENDER TABLE ─────────────────────────────────────────────
function rs(v) { return v ? '&#8377;' + v : '<span class="td-na">—</span>'; }
function rt(v) { return v || '<span class="td-na">—</span>'; }
function rgst(v, other) {
  if (v) return '<span class="td-gst">&#8377;' + v + '</span>';
  if (other) return '<span class="td-na">N/A</span>';
  return '<span class="td-na">—</span>';
}

function renderResults(rows, errors) {
  const blanks = rows.filter(r =>
    !r['Grand Total'] || (!r.IGST && !r.CGST && !r.SGST) || !r.Date
  ).length;
  document.getElementById('s-total').textContent = rows.length + (errors||[]).length;
  document.getElementById('s-ok').textContent    = rows.length;
  document.getElementById('s-err').textContent   = (errors||[]).length;
  document.getElementById('s-blank').textContent = blanks;
  renderTable(rows);
  const eb = document.getElementById('error-box');
  const el = document.getElementById('error-list');
  if (errors && errors.length) {
    el.innerHTML = errors.map(e =>
      `<div class="error-item"><strong>${e.File}</strong> &mdash; ${e.Error}</div>`
    ).join('');
    eb.style.display = 'block';
  } else {
    eb.style.display = 'none';
  }
  document.getElementById('results-section').style.display = 'block';
}

function renderTable(rows) {
  const tbody = document.getElementById('tbl-body');
  tbody.innerHTML = '';
  rows.forEach(r => {
    const hasI = !!r.IGST, hasC = !!(r.CGST || r.SGST);
    const tr = document.createElement('tr');
    tr.dataset.bill    = (r['Bill No'] || '').toLowerCase();
    tr.dataset.company = (r.Company || '').toLowerCase();
    tr.innerHTML =
      `<td class="td-bill">${rt(r['Bill No'])}</td>` +
      `<td class="td-company" title="${r.Company||''}">${rt(r.Company)}</td>` +
      `<td class="td-date">${rt(r.Date)}</td>` +
      `<td class="td-amt">${rs(r['Grand Total'])}</td>` +
      `<td>${rgst(r.IGST, hasC)}</td>` +
      `<td>${rgst(r.CGST, hasI)}</td>` +
      `<td>${rgst(r.SGST, hasI)}</td>`;
    tbody.appendChild(tr);
  });
}

function filterTable() {
  const q = document.getElementById('search-box').value.toLowerCase();
  document.querySelectorAll('#tbl-body tr').forEach(tr => {
    const match = tr.dataset.bill.includes(q) || tr.dataset.company.includes(q);
    tr.style.display = match ? '' : 'none';
  });
}

// ── DOWNLOAD ─────────────────────────────────────────────────
function downloadExcel() {
  if (sessionId) window.location.href = '/download/' + sessionId;
}

// ── RESET ────────────────────────────────────────────────────
function resetAll() {
  clearFile();
  document.getElementById('prog-wrap').classList.remove('show');
  document.getElementById('prog-fill').style.width = '0%';
  document.getElementById('prog-pct').textContent = '0%';
  document.getElementById('results-section').style.display = 'none';
  document.getElementById('nav-status').textContent = 'Ready';
  sessionId = null; allRows = [];
}

// ── HISTORY ──────────────────────────────────────────────────
function addToHistory(name, count, errors, sid) {
  history.unshift({
    id: sid,
    name: name,
    date: new Date().toLocaleString('en-IN', {day:'2-digit',month:'short',year:'numeric',hour:'2-digit',minute:'2-digit'}),
    count: count,
    errors: errors,
  });
}

function renderHistory() {
  const body = document.getElementById('history-body');
  if (!history.length) {
    body.innerHTML = `<div class="history-empty">
      <svg viewBox="0 0 24 24"><path d="M13 3c-4.97 0-9 4.03-9 9H1l3.89 3.89.07.14L9 12H6c0-3.87 3.13-7 7-7s7 3.13 7 7-3.13 7-7 7c-1.93 0-3.68-.79-4.94-2.06l-1.42 1.42C8.27 19.99 10.51 21 13 21c4.97 0 9-4.03 9-9s-4.03-9-9-9zm-1 5v5l4.28 2.54.72-1.21-3.5-2.08V8H12z"/></svg>
      <div style="font-weight:600;color:var(--navy);margin-bottom:.35rem">No history yet</div>
      <div style="font-size:.82rem">Your extraction history will appear here after you process bills.</div>
    </div>`;
    return;
  }
  body.innerHTML = `<div class="tbl-wrap" style="border-radius:0;border:none;">
    <table class="history-table">
      <thead><tr><th>File Name</th><th>Date &amp; Time</th><th>Bills</th><th>Status</th><th>Action</th></tr></thead>
      <tbody>
        ${history.map(h => `
          <tr>
            <td style="font-weight:600;color:var(--navy)">${h.name}</td>
            <td style="color:var(--muted);font-size:.8rem">${h.date}</td>
            <td><strong>${h.count}</strong> bills</td>
            <td>
              ${h.errors > 0
                ? `<span class="h-badge warn">${h.errors} errors</span>`
                : `<span class="h-badge ok">All read</span>`}
            </td>
            <td>
              <button class="btn btn-outline" style="padding:.35rem .8rem;font-size:.76rem"
                onclick="window.location.href='/download/${h.id}'">Download Excel</button>
            </td>
          </tr>`).join('')}
      </tbody>
    </table>
  </div>`;
}

function clearHistory() {
  if (!history.length) return;
  if (confirm('Clear all history? This cannot be undone.')) {
    history = [];
    renderHistory();
  }
}
</script>
</body>
</html>"""

# ============================================================
# ROUTES
# ============================================================

@app.route("/")
def index():
    return render_template_string(HTML)

@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    if not f.filename.endswith(".zip"):
        return jsonify({"error": "Please upload a .zip file"}), 400
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    f.save(tmp.name); tmp.close()
    results, errors = process_zip(tmp.name)
    os.unlink(tmp.name)
    sid = str(uuid.uuid4())
    SESSIONS[sid] = results
    # Save Excel to history folder
    try:
        buf = make_excel(results)
        fname = f"gst_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{sid[:8]}.xlsx"
        with open(os.path.join(HISTORY_DIR, fname), "wb") as fp:
            fp.write(buf.read())
    except Exception: pass
    return jsonify({"session_id": sid, "results": results, "errors": errors, "count": len(results)})

@app.route("/download/<session_id>")
def download(session_id):
    rows = SESSIONS.get(session_id)
    if rows is None: return "Session not found", 404
    buf = make_excel(rows)
    return send_file(buf, as_attachment=True, download_name="gst_details.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

if __name__ == "__main__":
    print("\n" + "="*50)
    print("  GST Manager  →  http://localhost:5000")
    print("="*50 + "\n")
    app.run(debug=True, port=5000)
